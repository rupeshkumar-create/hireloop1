"""
Resume parsing service — our own two-tier extraction (no paid third-party parser).

Parsing tiers (best result wins, lower tier augments missing fields):
  1. LLM (OpenRouter/Claude) — strong, schema-guided extraction from raw text:
     contact, full work history with dates, education, skills, location, links.
  2. Deterministic regex     — zero-network fallback, always available; also
     backfills anything the LLM missed and carries the result with no key at all.

Both tiers return the same `ParsedResume` model mapped to our candidates schema.
`parse_best()` runs the chain, merges results, and never raises — the worst case
is a sparse ParsedResume so onboarding still proceeds.
"""

# ruff: noqa: RUF001
# (en/em dashes are intentional here: resume date ranges and title/company
#  separators use unicode dashes, which we must match and strip literally.)

from __future__ import annotations

import io
import json
import re
from datetime import date, datetime

import httpx
import structlog
from pydantic import BaseModel, Field, ValidationError

from hireloop_api.services.skills import display_skill as _vocab_display
from hireloop_api.services.skills import is_known_skill as _vocab_known

logger = structlog.get_logger()

# ── Skill dictionary (canonical → aliases) ─────────────────────────────────────
# Detection matches any alias on a word boundary and records the canonical name.
# Kept broad but India-job-market relevant. The LLM tier finds the long tail;
# this is the deterministic safety net.
_SKILL_ALIASES: dict[str, list[str]] = {
    # Languages
    "python": ["python"],
    "java": ["java"],
    "javascript": ["javascript", "js", "es6"],
    "typescript": ["typescript", "ts"],
    "c++": ["c\\+\\+", "cpp"],
    "c#": ["c#", "c sharp", "csharp"],
    "go": ["golang", "go lang"],
    "rust": ["rust"],
    "ruby": ["ruby"],
    "php": ["php"],
    "kotlin": ["kotlin"],
    "swift": ["swift"],
    "scala": ["scala"],
    "r": ["\\br language\\b"],
    "sql": ["sql"],
    "bash": ["bash", "shell scripting"],
    # Frontend
    "react": ["react", "react.js", "reactjs"],
    "next.js": ["next.js", "nextjs"],
    "vue": ["vue", "vue.js", "vuejs"],
    "angular": ["angular", "angularjs"],
    "redux": ["redux"],
    "tailwind": ["tailwind", "tailwindcss"],
    "html": ["html", "html5"],
    "css": ["css", "css3", "sass", "scss"],
    # Backend / frameworks
    "node.js": ["node.js", "nodejs", "node js", "node"],
    "express": ["express", "express.js"],
    "fastapi": ["fastapi"],
    "django": ["django"],
    "flask": ["flask"],
    "spring": ["spring", "spring boot", "springboot"],
    "rails": ["rails", "ruby on rails"],
    "dotnet": [".net", "dotnet", "asp.net"],
    "graphql": ["graphql"],
    "rest": ["rest api", "restful", "rest"],
    "grpc": ["grpc"],
    # Databases
    "postgresql": ["postgresql", "postgres"],
    "mysql": ["mysql"],
    "mongodb": ["mongodb", "mongo"],
    "redis": ["redis"],
    "elasticsearch": ["elasticsearch", "elastic search"],
    "dynamodb": ["dynamodb"],
    "cassandra": ["cassandra"],
    "sqlite": ["sqlite"],
    "snowflake": ["snowflake"],
    "bigquery": ["bigquery", "big query"],
    "pgvector": ["pgvector"],
    # Cloud / DevOps
    "aws": ["aws", "amazon web services"],
    "gcp": ["gcp", "google cloud"],
    "azure": ["azure"],
    "docker": ["docker"],
    "kubernetes": ["kubernetes", "k8s"],
    "terraform": ["terraform"],
    "jenkins": ["jenkins"],
    "github actions": ["github actions"],
    "ci/cd": ["ci/cd", "cicd", "ci cd"],
    "linux": ["linux", "unix"],
    "nginx": ["nginx"],
    "kafka": ["kafka"],
    "rabbitmq": ["rabbitmq"],
    "airflow": ["airflow"],
    # Data / ML / AI
    "machine learning": ["machine learning", "\\bml\\b"],
    "deep learning": ["deep learning"],
    "data science": ["data science"],
    "data analysis": ["data analysis", "data analytics"],
    "nlp": ["nlp", "natural language processing"],
    "computer vision": ["computer vision", "opencv"],
    "pytorch": ["pytorch"],
    "tensorflow": ["tensorflow"],
    "scikit-learn": ["scikit-learn", "sklearn"],
    "pandas": ["pandas"],
    "numpy": ["numpy"],
    "spark": ["spark", "pyspark"],
    "hadoop": ["hadoop"],
    "tableau": ["tableau"],
    "power bi": ["power bi", "powerbi"],
    "excel": ["excel", "advanced excel"],
    "llm": ["llm", "large language model", "gpt", "openai"],
    "langchain": ["langchain"],
    "langgraph": ["langgraph"],
    "rag": ["rag", "retrieval augmented"],
    # Mobile
    "android": ["android"],
    "ios": ["ios"],
    "flutter": ["flutter"],
    "react native": ["react native"],
    # Design
    "figma": ["figma"],
    "sketch": ["sketch"],
    "adobe xd": ["adobe xd"],
    "photoshop": ["photoshop"],
    "illustrator": ["illustrator"],
    "ui/ux": ["ui/ux", "ux design", "ui design", "user experience"],
    # Business / role skills
    "product management": ["product management", "product manager"],
    "project management": ["project management", "pmp"],
    "agile": ["agile", "scrum", "kanban"],
    "jira": ["jira"],
    "sales": ["sales", "business development", "bd"],
    "marketing": ["marketing", "digital marketing"],
    "seo": ["seo", "search engine optimization"],
    "content writing": ["content writing", "copywriting"],
    "customer success": ["customer success", "account management"],
    "finance": ["finance", "financial analysis"],
    "accounting": ["accounting", "tally"],
    "hr": ["human resources", "recruitment", "talent acquisition"],
    "operations": ["operations management", "supply chain"],
    "communication": ["communication skills"],
    "leadership": ["leadership", "team management"],
}

_KNOWN_SKILLS = set(_SKILL_ALIASES.keys())

# Canonical skill keys are lowercase for matching; this maps them to the casing
# people actually expect to see on a profile. Anything not listed is titleized.
_SKILL_DISPLAY: dict[str, str] = {
    "javascript": "JavaScript",
    "typescript": "TypeScript",
    "c++": "C++",
    "c#": "C#",
    "php": "PHP",
    "sql": "SQL",
    "html": "HTML",
    "css": "CSS",
    "react": "React",
    "next.js": "Next.js",
    "vue": "Vue",
    "node.js": "Node.js",
    "express": "Express",
    "fastapi": "FastAPI",
    "graphql": "GraphQL",
    "grpc": "gRPC",
    "rest": "REST",
    "postgresql": "PostgreSQL",
    "mysql": "MySQL",
    "mongodb": "MongoDB",
    "dynamodb": "DynamoDB",
    "sqlite": "SQLite",
    "bigquery": "BigQuery",
    "dotnet": ".NET",
    "aws": "AWS",
    "gcp": "GCP",
    "ios": "iOS",
    "ml": "Machine Learning",
    "ai": "AI",
    "nlp": "NLP",
    "ci/cd": "CI/CD",
    "devops": "DevOps",
    "ui/ux": "UI/UX",
    "pytorch": "PyTorch",
    "tensorflow": "TensorFlow",
    "scikit-learn": "scikit-learn",
    "numpy": "NumPy",
    "pandas": "pandas",
    "kubernetes": "Kubernetes",
    "github": "GitHub",
    "gitlab": "GitLab",
    "mlops": "MLOps",
    "nosql": "NoSQL",
    "saas": "SaaS",
    "api": "API",
}


def _display_skill(canonical: str) -> str:
    """Return human-friendly casing for a canonical (lowercase) skill key."""
    if canonical in _SKILL_DISPLAY:
        return _SKILL_DISPLAY[canonical]
    if canonical in _KNOWN_SKILLS:
        return canonical.title()
    # Free-text skill from an explicit "Skills:" section — keep as-is if it
    # already has mixed case, otherwise titleize a plain lowercase token.
    return canonical if canonical != canonical.lower() else canonical.title()


# Indian states + UTs and their major cities → state mapping.
_INDIAN_LOCATIONS: dict[str, str] = {
    "bengaluru": "Karnataka",
    "bangalore": "Karnataka",
    "mysore": "Karnataka",
    "mysuru": "Karnataka",
    "mumbai": "Maharashtra",
    "pune": "Maharashtra",
    "nagpur": "Maharashtra",
    "nashik": "Maharashtra",
    "thane": "Maharashtra",
    "navi mumbai": "Maharashtra",
    "delhi": "Delhi",
    "new delhi": "Delhi",
    "gurugram": "Haryana",
    "gurgaon": "Haryana",
    "faridabad": "Haryana",
    "noida": "Uttar Pradesh",
    "ghaziabad": "Uttar Pradesh",
    "lucknow": "Uttar Pradesh",
    "kanpur": "Uttar Pradesh",
    "hyderabad": "Telangana",
    "secunderabad": "Telangana",
    "chennai": "Tamil Nadu",
    "coimbatore": "Tamil Nadu",
    "madurai": "Tamil Nadu",
    "kolkata": "West Bengal",
    "howrah": "West Bengal",
    "ahmedabad": "Gujarat",
    "surat": "Gujarat",
    "vadodara": "Gujarat",
    "gandhinagar": "Gujarat",
    "rajkot": "Gujarat",
    "jaipur": "Rajasthan",
    "jodhpur": "Rajasthan",
    "udaipur": "Rajasthan",
    "indore": "Madhya Pradesh",
    "bhopal": "Madhya Pradesh",
    "kochi": "Kerala",
    "cochin": "Kerala",
    "thiruvananthapuram": "Kerala",
    "trivandrum": "Kerala",
    "kozhikode": "Kerala",
    "bhubaneswar": "Odisha",
    "patna": "Bihar",
    "ranchi": "Jharkhand",
    "raipur": "Chhattisgarh",
    "chandigarh": "Chandigarh",
    "dehradun": "Uttarakhand",
    "guwahati": "Assam",
    "visakhapatnam": "Andhra Pradesh",
    "vijayawada": "Andhra Pradesh",
    "amritsar": "Punjab",
    "ludhiana": "Punjab",
    "goa": "Goa",
    "panaji": "Goa",
}

_INDIAN_STATE_ALIASES: dict[str, str] = {
    "ka": "Karnataka",
    "karnataka": "Karnataka",
    "mh": "Maharashtra",
    "maharashtra": "Maharashtra",
    "dl": "Delhi",
    "delhi": "Delhi",
    "tn": "Tamil Nadu",
    "tamil nadu": "Tamil Nadu",
    "ts": "Telangana",
    "tg": "Telangana",
    "telangana": "Telangana",
    "up": "Uttar Pradesh",
    "uttar pradesh": "Uttar Pradesh",
    "hr": "Haryana",
    "haryana": "Haryana",
    "gj": "Gujarat",
    "gujarat": "Gujarat",
    "rj": "Rajasthan",
    "rajasthan": "Rajasthan",
    "mp": "Madhya Pradesh",
    "madhya pradesh": "Madhya Pradesh",
    "kl": "Kerala",
    "kerala": "Kerala",
    "ap": "Andhra Pradesh",
    "andhra pradesh": "Andhra Pradesh",
    "pb": "Punjab",
    "punjab": "Punjab",
}

_CITY_DISPLAY_OVERRIDES = {
    "bangalore": "Bengaluru",
    "bengaluru": "Bengaluru",
    "cochin": "Kochi",
    "trivandrum": "Thiruvananthapuram",
}

_JUNK_SKILLS = {
    "team player",
    "hard working",
    "hardworking",
    "self motivated",
    "self-motivated",
    "quick learner",
    "responsible",
    "punctual",
    "honest",
    "dedicated",
}

# Lines that should never be treated as a candidate name.
_SECTION_HEADERS = {
    "resume",
    "curriculum vitae",
    "cv",
    "profile",
    "summary",
    "objective",
    "experience",
    "work experience",
    "professional experience",
    "education",
    "skills",
    "technical skills",
    "projects",
    "certifications",
    "contact",
    "personal details",
}

_DATE_RANGE = re.compile(
    r"((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*'?\d{2,4}"
    r"|\d{1,2}/\d{4}|\d{4})\s*[-–—to]+\s*"
    r"((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*'?\d{2,4}"
    r"|\d{1,2}/\d{4}|\d{4}|present|current|now|till date|ongoing)",
    re.IGNORECASE,
)


class WorkExperience(BaseModel):
    company: str | None = None
    title: str | None = None
    start_date: str | None = None
    end_date: str | None = None
    description: str | None = None
    is_current: bool = False


class Education(BaseModel):
    institution: str | None = None
    degree: str | None = None
    field_of_study: str | None = None
    start_date: str | None = None
    end_date: str | None = None
    grade: str | None = None


class ParsedResume(BaseModel):
    """Structured data extracted from a resume."""

    full_name: str | None = None
    email: str | None = None
    phone: str | None = None
    headline: str | None = None
    summary: str | None = None
    current_title: str | None = None
    current_company: str | None = None
    years_experience: int | None = None
    # Indian-résumé staples — extracting these means the profile form / Aarya
    # don't have to re-ask them. CTC values are INR per annum; notice is in days.
    expected_ctc_min: int | None = None
    expected_ctc_max: int | None = None
    current_ctc: int | None = None
    notice_period_days: int | None = None
    skills: list[str] = []
    work_experience: list[WorkExperience] = []
    education: list[Education] = []
    linkedin_url: str | None = None
    github_url: str | None = None
    location_city: str | None = None
    location_state: str | None = None
    career_profile: dict = Field(default_factory=dict)
    career_analysis: dict = Field(default_factory=dict)
    parser_metadata: dict = Field(default_factory=dict)
    raw_text: str | None = None


# ── LLM extraction prompt ──────────────────────────────────────────────────────

_LLM_SYSTEM_PROMPT = """You are an expert resume parser for an Indian recruiting platform.
Extract structured data from the resume text and return ONLY valid JSON (no prose,
no markdown fences) matching exactly this shape:

{
  "full_name": string|null,
  "email": string|null,
  "phone": string|null,
  "headline": string|null,            // a short professional headline
  "summary": string|null,             // 1-3 sentence professional summary
  "current_title": string|null,
  "current_company": string|null,
  "years_experience": number|null,    // total full-time years, integer
  "current_ctc": number|null,         // INR per annum, integer (e.g. "18 LPA" → 1800000)
  "expected_ctc_min": number|null,    // INR per annum, integer
  "expected_ctc_max": number|null,    // INR per annum, integer
  "notice_period_days": number|null,  // integer days ("2 months" → 60, "Immediate" → 0)
  "skills": [string],                 // lowercase, deduped, specific (max 40)
  "work_experience": [
    {"company": string|null, "title": string|null,
     "start_date": string|null,       // "YYYY-MM" or "YYYY"
     "end_date": string|null,         // "YYYY-MM", "YYYY", or null if current
     "description": string|null, "is_current": boolean}
  ],
  "education": [
    {"institution": string|null, "degree": string|null, "field_of_study": string|null,
     "start_date": string|null, "end_date": string|null, "grade": string|null}
  ],
  "linkedin_url": string|null,
  "github_url": string|null,
  "location_city": string|null,       // Indian city if present
  "location_state": string|null       // Indian state
}

Rules:
- Use null for anything genuinely absent. Never invent data.
- Order work_experience most-recent first; mark the current role is_current=true.
- For Indian phone numbers keep the +91 / leading-zero form as written.
- Skills must be concrete technologies/competencies, not full sentences.
- Derive skills from what the person actually DID in their work_experience and
  projects. IGNORE a LinkedIn "Top Skills" sidebar or any standalone skills list
  that the experience does not support — those are frequently stale/aspirational
  and misrepresent the candidate (e.g. "Sales Operations" listed by a UX
  designer). When a sidebar skill conflicts with the actual roles, omit it."""


class ResumeParserService:
    """Our own two-tier resume parser: LLM (schema-guided) → deterministic regex,
    with field-level merge. No third-party paid parser — the LLM tier extracts the
    full structured profile; regex is the always-available, zero-network fallback."""

    OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"

    # ── Public orchestrator ─────────────────────────────────────────────────

    @classmethod
    async def parse_best(
        cls,
        *,
        file_bytes: bytes,
        filename: str,
        mime_type: str | None,
        settings: object | None = None,
    ) -> ParsedResume:
        """
        Run the full parsing chain and return the richest possible result.

        Never raises: any tier failure degrades to the next tier, and the final
        regex tier always produces a (possibly sparse) ParsedResume.
        """
        text = cls._extract_text(file_bytes, filename, mime_type)
        regex_result = cls.parse_from_text(text)

        openrouter_key = getattr(settings, "openrouter_api_key", "") if settings else ""
        model = (
            getattr(settings, "openrouter_primary_model", "") or "anthropic/claude-opus-4.7"
            if settings
            else "anthropic/claude-opus-4.7"
        )

        best: ParsedResume | None = None

        # Tier 1 (primary): our LLM parser — strong, schema-guided, extracts the
        # full structured profile (contact, work history with dates, education,
        # skills, location, links).
        if openrouter_key and text.strip():
            try:
                llm_result = await cls._parse_with_llm(
                    text=text, api_key=openrouter_key, model=model
                )
                if llm_result:
                    logger.info("resume_parsed_via", tier="llm")
                    best = llm_result
            except Exception as exc:
                logger.warning("llm_tier_failed", error=str(exc)[:200])

        # Tier 1b (vision fallback): a scanned/image-only PDF yields almost no
        # text, so the text tiers can't help. Render the pages and let Claude's
        # vision model read them. No OCR engine, no new vendor — reuses the LLM.
        is_pdf = (mime_type == "application/pdf") or (filename or "").lower().endswith(".pdf")
        scanned = is_pdf and len(text.strip()) < cls.SCAN_TEXT_THRESHOLD
        if best is None and openrouter_key and scanned:
            try:
                vision_result = await cls._parse_with_vision(
                    file_bytes=file_bytes, api_key=openrouter_key, model=model
                )
                if vision_result:
                    logger.info("resume_parsed_via", tier="vision")
                    best = vision_result
            except Exception as exc:
                logger.warning("vision_tier_failed", error=str(exc)[:200])

        # Tier 2 (fallback / augmentation): deterministic regex — always available,
        # backfills anything the LLM missed (or carries the whole result with no key).
        best = _merge(best, regex_result) if best else regex_result

        # Always retain extracted text for downstream embedding/search.
        if not best.raw_text:
            best.raw_text = text or None
        _normalise_parsed_resume(best, source="best")
        _ensure_career_profile(best)
        return best

    @staticmethod
    def parse_from_bytes_local(
        file_bytes: bytes, filename: str, mime_type: str | None
    ) -> ParsedResume:
        """Best-effort local parser (regex) used when remote tiers are unavailable."""
        text = ResumeParserService._extract_text(file_bytes, filename, mime_type)
        return ResumeParserService.parse_from_text(text)

    # ── LLM tier ─────────────────────────────────────────────────────────────

    @classmethod
    async def _parse_with_llm(cls, *, text: str, api_key: str, model: str) -> ParsedResume | None:
        """Extract structured fields via OpenRouter (Claude). Returns None on failure."""
        # Keep prompt bounded — resumes rarely need more than this much text.
        snippet = text.strip()[:16000]
        payload = {
            "model": model,
            "temperature": 0,
            "max_tokens": 2000,
            "response_format": {"type": "json_object"},
            "messages": [
                {"role": "system", "content": _LLM_SYSTEM_PROMPT},
                {"role": "user", "content": f"Resume text:\n\n{snippet}"},
            ],
        }
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://app.hireloop.in",
            "X-Title": "Hireloop - Resume Parser",
        }
        async with httpx.AsyncClient(timeout=45.0) as client:
            resp = await client.post(cls.OPENROUTER_URL, headers=headers, json=payload)
        if resp.status_code != 200:
            logger.warning("llm_parse_http_error", status=resp.status_code, body=resp.text[:200])
            return None

        try:
            content = resp.json()["choices"][0]["message"]["content"]
        except (KeyError, IndexError, ValueError):
            return None

        data = _loads_json_lenient(content)
        if not isinstance(data, dict):
            return None
        return cls._parsed_from_llm_data(data, text=text, source="llm")

    @classmethod
    def _parsed_from_llm_data(cls, data: dict, *, text: str, source: str) -> ParsedResume | None:
        """Build a ParsedResume from the model's JSON (shared by the text-LLM and
        the vision tiers). Returns None if the payload fails validation."""
        try:
            work = [
                WorkExperience(**w) for w in data.get("work_experience", []) if isinstance(w, dict)
            ]
            edu = [Education(**e) for e in data.get("education", []) if isinstance(e, dict)]
            skills = _clean_skill_list(data.get("skills", []))
            years = data.get("years_experience")
            years_int = (
                int(years) if isinstance(years, (int, float)) else _infer_years_experience(text)
            )
            current = next((w for w in work if w.is_current), work[0] if work else None)
            parsed = ParsedResume(
                full_name=_clean_str(data.get("full_name")),
                email=_clean_str(data.get("email")),
                phone=_clean_str(data.get("phone")),
                headline=(_clean_str(data.get("headline")) or (current.title if current else None)),
                summary=_clean_str(data.get("summary")),
                current_title=(
                    _clean_str(data.get("current_title"))
                    or (current.title if current else None)
                    or _first_work_title(work)
                ),
                current_company=(
                    _clean_str(data.get("current_company"))
                    or (current.company if current else None)
                ),
                years_experience=min(60, years_int) if years_int is not None else None,
                current_ctc=_clean_int(data.get("current_ctc")),
                expected_ctc_min=_clean_int(data.get("expected_ctc_min")),
                expected_ctc_max=_clean_int(data.get("expected_ctc_max")),
                notice_period_days=_clean_int(data.get("notice_period_days")),
                skills=skills,
                work_experience=work,
                education=edu,
                linkedin_url=_clean_str(data.get("linkedin_url")),
                github_url=_clean_str(data.get("github_url")),
                location_city=_clean_str(data.get("location_city")),
                location_state=_clean_str(data.get("location_state")),
                raw_text=text or None,
            )
            _normalise_parsed_resume(parsed, source=source)
            _ensure_career_profile(parsed)
            return parsed
        except ValidationError as exc:
            logger.warning("llm_parse_validation_failed", error=str(exc)[:200], tier=source)
            return None

    # ── Vision tier (scanned / image-only PDFs) ────────────────────────────────

    # Below this many extracted characters a PDF is almost certainly a scan/image
    # (no embedded text layer) — pdfplumber/pypdf can't help, so we fall back to
    # Claude's vision model on rendered page images.
    SCAN_TEXT_THRESHOLD = 120

    @staticmethod
    def _render_pdf_to_images(
        file_bytes: bytes, *, max_pages: int = 3, scale: float = 2.0
    ) -> list[bytes]:
        """Render the first few PDF pages to PNG bytes via pypdfium2 (already a
        pdfplumber dependency; permissive licence, no system binaries). Returns
        [] on any failure so the caller degrades gracefully."""
        try:
            import pypdfium2 as pdfium

            images: list[bytes] = []
            pdf = pdfium.PdfDocument(file_bytes)
            try:
                for i in range(min(len(pdf), max_pages)):
                    bitmap = pdf[i].render(scale=scale)
                    pil_image = bitmap.to_pil()
                    buf = io.BytesIO()
                    pil_image.save(buf, format="PNG")
                    images.append(buf.getvalue())
            finally:
                pdf.close()
            return images
        except Exception as exc:
            logger.warning("pdf_render_failed", error=str(exc)[:200])
            return []

    @classmethod
    async def _parse_with_vision(
        cls, *, file_bytes: bytes, api_key: str, model: str
    ) -> ParsedResume | None:
        """Last-resort parser for scanned/image PDFs: render pages to images and
        ask Claude (vision) for the same structured profile. Returns None on
        failure. No OCR engine and no third-party parser — reuses the LLM we
        already pay for."""
        import base64

        images = cls._render_pdf_to_images(file_bytes)
        if not images:
            return None

        content: list[dict] = [
            {
                "type": "text",
                "text": (
                    "This is a scanned resume (image only). Read it and return "
                    "the structured JSON profile."
                ),
            }
        ]
        for png in images:
            b64 = base64.b64encode(png).decode("ascii")
            content.append(
                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}}
            )

        payload = {
            "model": model,
            "temperature": 0,
            "max_tokens": 2000,
            "response_format": {"type": "json_object"},
            "messages": [
                {"role": "system", "content": _LLM_SYSTEM_PROMPT},
                {"role": "user", "content": content},
            ],
        }
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://app.hireloop.in",
            "X-Title": "Hireloop - Resume Parser (vision)",
        }
        try:
            async with httpx.AsyncClient(timeout=60.0) as client:
                resp = await client.post(cls.OPENROUTER_URL, headers=headers, json=payload)
        except Exception as exc:
            logger.warning("vision_parse_request_failed", error=str(exc)[:200])
            return None
        if resp.status_code != 200:
            logger.warning("vision_parse_http_error", status=resp.status_code, body=resp.text[:200])
            return None
        try:
            msg = resp.json()["choices"][0]["message"]["content"]
        except (KeyError, IndexError, ValueError):
            return None
        data = _loads_json_lenient(msg)
        if not isinstance(data, dict):
            return None
        return cls._parsed_from_llm_data(data, text="", source="vision")

    # ── Text extraction ───────────────────────────────────────────────────────

    @staticmethod
    def _extract_text(file_bytes: bytes, filename: str, mime_type: str | None) -> str:
        lowered = (filename or "").lower()
        if mime_type == "application/pdf" or lowered.endswith(".pdf"):
            return ResumeParserService._extract_pdf_text(file_bytes)
        if (
            mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or lowered.endswith(".docx")
        ):
            return ResumeParserService._extract_docx_text(file_bytes)
        # Unknown type: try PDF then DOCX then treat as utf-8 text.
        text = ResumeParserService._extract_pdf_text(file_bytes)
        if text:
            return text
        text = ResumeParserService._extract_docx_text(file_bytes)
        if text:
            return text
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            return ""

    @staticmethod
    def _extract_pdf_text(file_bytes: bytes) -> str:
        """
        Layout-aware extraction (#16): two-column CVs read column-by-column
        instead of interleaving left/right lines (which scrambles work history
        for the LLM tier). pdfplumber gives word positions; when a page's words
        cluster into two x-bands separated by a gutter, each band is extracted
        as its own column in reading order. Falls back to pypdf on any failure.
        """
        try:
            import pdfplumber

            pages: list[str] = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    words = page.extract_words() or []
                    midpoint = page.width / 2
                    # A real gutter means almost no word CROSSES the midline,
                    # and both halves hold a meaningful share of the words.
                    crossing = sum(1 for w in words if w["x0"] < midpoint < w["x1"])
                    left = [w for w in words if w["x1"] <= midpoint]
                    right = [w for w in words if w["x0"] >= midpoint]
                    two_col = (
                        len(words) >= 40
                        and crossing / len(words) < 0.04
                        and min(len(left), len(right)) / len(words) > 0.2
                    )
                    if two_col:
                        col_texts = []
                        for half in (left, right):
                            col_texts.append(_words_to_lines(half))
                        pages.append("\n\n".join(col_texts))
                    else:
                        pages.append(page.extract_text() or "")
            text = "\n".join(pages)
            if text.strip():
                return text
        except Exception as exc:
            logger.warning("pdfplumber_parse_failed", error=str(exc))
        # Fallback: plain pypdf (previous behavior).
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(file_bytes))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as exc:
            logger.warning("local_pdf_parse_failed", error=str(exc))
            return ""

    @staticmethod
    def _extract_docx_text(file_bytes: bytes) -> str:
        try:
            from docx import Document

            document = Document(io.BytesIO(file_bytes))
            parts = [p.text for p in document.paragraphs]
            # Tables often hold skills/contact grids — pull those too.
            for table in document.tables:
                for row in table.rows:
                    parts.append(" ".join(cell.text for cell in row.cells))
            return "\n".join(parts)
        except Exception as exc:
            logger.warning("local_docx_parse_failed", error=str(exc))
            return ""

    # ── Regex / heuristic tier ──────────────────────────────────────────────

    @staticmethod
    def parse_from_text(text: str) -> ParsedResume:
        """Infer candidate fields from plain resume text (deterministic)."""
        clean = _normalise_text(text)
        if not clean.strip():
            return ParsedResume()
        lines = [line.strip() for line in clean.splitlines() if line.strip()]

        full_name = _infer_name(lines)
        email = _first_match(clean, r"[\w.+-]+@[\w-]+\.[\w.-]+")
        phone = _infer_india_phone(clean)
        linkedin_url = _infer_profile_url(clean, "linkedin")
        github_url = _infer_profile_url(clean, "github")

        work_exp = _infer_work_experience(lines)
        title, company = _infer_current_role(lines)
        if work_exp:
            current = next((w for w in work_exp if w.is_current), work_exp[0])
            title = current.title or title
            company = current.company or company

        years_experience = _infer_years_experience(clean) or _estimate_years_static(work_exp)
        skills = _infer_skills(clean)
        city, state = _infer_location(clean)
        education = _infer_education(lines)
        summary = _infer_summary(lines, title, company, years_experience, skills)
        exp_ctc_min, exp_ctc_max, cur_ctc = _infer_ctc(clean)
        notice_days = _infer_notice_period(clean)

        if not work_exp and (title or company):
            work_exp = [WorkExperience(company=company, title=title, is_current=True)]

        parsed = ParsedResume(
            full_name=full_name,
            email=email,
            phone=phone,
            headline=title,
            summary=summary,
            current_title=title,
            current_company=company,
            years_experience=years_experience,
            expected_ctc_min=exp_ctc_min,
            expected_ctc_max=exp_ctc_max,
            current_ctc=cur_ctc,
            notice_period_days=notice_days,
            skills=skills,
            work_experience=work_exp,
            education=education,
            linkedin_url=linkedin_url,
            github_url=github_url,
            location_city=city,
            location_state=state,
            raw_text=clean or None,
        )
        _normalise_parsed_resume(parsed, source="regex")
        _ensure_career_profile(parsed)
        return parsed

    @staticmethod
    def _estimate_years_experience(work_exp: list[WorkExperience]) -> int | None:
        if not work_exp:
            return None
        total_months = 0
        for exp in work_exp:
            if exp.start_date:
                try:
                    start = datetime.strptime(exp.start_date[:7], "%Y-%m").date()
                    end = (
                        datetime.strptime(exp.end_date[:7], "%Y-%m").date()
                        if exp.end_date
                        else date.today()
                    )
                    total_months += (end.year - start.year) * 12 + (end.month - start.month)
                except ValueError:
                    total_months += 24
            else:
                total_months += 24
        return max(0, round(total_months / 12))


# ── Layout helpers (#16) ───────────────────────────────────────────────────────


def _words_to_lines(words: list[dict], line_tolerance: float = 3.0) -> str:
    """
    Rebuild reading-order text from pdfplumber word boxes: group words whose
    vertical position is within `line_tolerance` pts into one line (top→bottom),
    join each line left→right. Used per column on two-column pages.
    """
    if not words:
        return ""
    ordered = sorted(words, key=lambda w: (w["top"], w["x0"]))
    lines: list[list[dict]] = [[ordered[0]]]
    for w in ordered[1:]:
        if abs(w["top"] - lines[-1][0]["top"]) <= line_tolerance:
            lines[-1].append(w)
        else:
            lines.append([w])
    return "\n".join(
        " ".join(w["text"] for w in sorted(line, key=lambda w: w["x0"])) for line in lines
    )


# ── Field-level merge ──────────────────────────────────────────────────────────


def _merge(primary: ParsedResume | None, secondary: ParsedResume | None) -> ParsedResume:
    """Fill empty fields in `primary` from `secondary`. Skills are unioned."""
    if primary is None:
        return secondary or ParsedResume()
    if secondary is None:
        return primary

    scalar_fields = [
        "full_name",
        "email",
        "phone",
        "headline",
        "summary",
        "current_title",
        "current_company",
        "years_experience",
        "expected_ctc_min",
        "expected_ctc_max",
        "current_ctc",
        "notice_period_days",
        "linkedin_url",
        "github_url",
        "location_city",
        "location_state",
        "raw_text",
    ]
    for field in scalar_fields:
        if not getattr(primary, field):
            val = getattr(secondary, field)
            if val:
                setattr(primary, field, val)

    # Union skills, preserve order, dedupe.
    merged_skills: list[str] = []
    for skill in [*primary.skills, *secondary.skills]:
        if skill and skill not in merged_skills:
            merged_skills.append(skill)
    primary.skills = merged_skills[:40]

    # Prefer the richer list for nested structures.
    if len(secondary.work_experience) > len(primary.work_experience):
        primary.work_experience = secondary.work_experience
    if len(secondary.education) > len(primary.education):
        primary.education = secondary.education
    _ensure_career_profile(primary)
    return primary


# ── Structured career profile builder ─────────────────────────────────────────


def _ensure_career_profile(parsed: ParsedResume) -> None:
    if parsed.career_profile and parsed.career_analysis:
        return
    profile = _build_career_profile(parsed)
    parsed.career_profile = parsed.career_profile or profile
    parsed.career_analysis = (
        parsed.career_analysis
        or profile["aspirations_market_fit_recommendations"]["career_path_recommendation"]
    )


def _build_career_profile(parsed: ParsedResume) -> dict:
    roles = [_work_experience_to_profile(role) for role in parsed.work_experience]
    total_experience = parsed.years_experience or _estimate_years_static(parsed.work_experience)
    job_changes = max(0, len(parsed.work_experience) - 1)
    seniority = _infer_seniority_level(parsed.current_title, total_experience)
    career_track = _infer_career_track(parsed.current_title, parsed.skills)
    industry = _infer_primary_industry(parsed.raw_text or "")
    hard_skills, soft_skills = _split_skill_groups(parsed.skills)
    next_roles = _next_roles(parsed.current_title, career_track, seniority)
    gaps = _gap_analysis(career_track, parsed.skills)

    profile_demographics = {
        "full_name": parsed.full_name,
        "current_location": _join_location(parsed.location_city, parsed.location_state),
        "preferred_work_location": None,
        "nationality_work_authorization": None,
        "years_of_experience": total_experience,
        "current_employment_status": "employed" if parsed.current_company else None,
        "contact_information": {
            "email": parsed.email,
            "phone": parsed.phone,
            "linkedin_url": parsed.linkedin_url,
            "github_url": parsed.github_url,
        },
        "languages_spoken": _infer_languages(parsed.raw_text or ""),
    }

    experience_history = {
        "current_career_snapshot": {
            "current_job_title": parsed.current_title,
            "current_company": parsed.current_company,
            "current_department_function": _infer_function(parsed.current_title),
            "current_seniority_level": seniority,
            "industry": industry,
            "career_track": career_track,
        },
        "roles": roles,
        "derived_metrics": {
            "total_experience": total_experience,
            "average_tenure": _average_tenure(parsed.work_experience),
            "number_of_job_changes": job_changes,
            "internal_promotions": _infer_promotions(parsed.work_experience),
            "career_stability_score": _career_stability_score(parsed.work_experience),
        },
    }

    skills_competencies = {
        "hard_skills": hard_skills,
        "soft_skills": soft_skills,
        "emerging_skills": _emerging_skills(parsed.skills),
        "skill_proficiency": [{"skill": skill, "level": "intermediate"} for skill in parsed.skills],
    }

    education_credentials = {
        "education": [
            {
                "degree": edu.degree,
                "major": edu.field_of_study,
                "minor": None,
                "university": edu.institution,
                "graduation_year": _year_from_date(edu.end_date),
                "gpa": edu.grade,
                "certifications_during_education": [],
                "academic_honors": [],
            }
            for edu in parsed.education
        ],
        "certifications": _infer_certifications(parsed.raw_text or ""),
    }

    achievements_leadership = {
        "achievements_impact": _infer_achievements(parsed.raw_text or ""),
        "leadership_experience": {
            "team_management_experience": _has_any(parsed.raw_text, ["managed", "led", "lead"]),
            "direct_reports": _first_int_after(parsed.raw_text or "", ["managed", "team of"]),
            "cross_functional_leadership": _has_any(
                parsed.raw_text, ["cross-functional", "stakeholders", "stakeholder"]
            ),
            "budget_ownership": _has_any(parsed.raw_text, ["budget", "portfolio"]),
            "hiring_experience": _has_any(parsed.raw_text, ["hiring", "recruiting", "interviewed"]),
            "mentoring_experience": _has_any(parsed.raw_text, ["mentor", "mentored", "coached"]),
            "executive_exposure": _has_any(
                parsed.raw_text, ["executive", "leadership team", "cxo"]
            ),
        },
        "industry_expertise": {
            "primary_industry": industry,
            "secondary_industries": [],
            "industry_depth": "moderate" if industry else None,
        },
        "domain_expertise": _infer_domain_expertise(parsed.current_title, parsed.skills),
        "reputation_signals": {
            "recommendations_received": None,
            "endorsements": [],
            "awards": _infer_awards(parsed.raw_text or ""),
            "publications": [],
            "patents": [],
            "conference_speaking": [],
            "open_source_contributions": [],
        },
    }

    career_recommendation = {
        "current_position": parsed.current_title,
        "current_level": seniority,
        "current_career_track": career_track,
        "next_likely_roles_1_3_years": next_roles[:3],
        "medium_term_roles_3_5_years": _medium_term_roles(career_track),
        "long_term_roles_5_10_years": _long_term_roles(career_track),
        "gap_analysis": gaps,
        "recommended_actions": _recommended_actions(gaps),
        "market_readiness": {
            "employability_score": _employability_score(parsed),
            "market_competitiveness_score": _market_competitiveness_score(parsed),
        },
        "career_progression_analysis": {
            "promotion_velocity": None,
            "title_growth": [role["job_title"] for role in roles if role["job_title"]],
            "responsibility_growth": None,
            "leadership_growth": None,
            "compensation_growth": None,
        },
    }

    return {
        "profile_demographics": profile_demographics,
        "experience_career_history": experience_history,
        "skills_competencies": skills_competencies,
        "education_credentials": education_credentials,
        "achievements_leadership": achievements_leadership,
        "aspirations_market_fit_recommendations": {
            "career_interests": _infer_career_interests(parsed.raw_text or "", parsed.skills),
            "career_transition_indicators": _infer_transition_indicators(parsed),
            "learning_growth_signals": _infer_learning_growth(parsed.raw_text or ""),
            "career_goals": {
                "desired_role": None,
                "desired_industry": None,
                "desired_seniority": None,
                "salary_expectations": None,
                "geographic_preferences": _join_location(
                    parsed.location_city, parsed.location_state
                ),
            },
            "career_path_recommendation": career_recommendation,
        },
    }


def _work_experience_to_profile(role: WorkExperience) -> dict:
    return {
        "job_title": role.title,
        "company": role.company,
        "employment_type": None,
        "start_date": role.start_date,
        "end_date": role.end_date,
        "duration": _duration_label(role.start_date, role.end_date),
        "industry": None,
        "team_size_managed": None,
        "reporting_structure": None,
        "responsibilities": _sentences(role.description),
        "achievements": _infer_achievements(role.description or ""),
        "promotions": [],
        "revenue_impact": _impact_metric(role.description or "", ["revenue", "sales"]),
        "cost_savings": _impact_metric(role.description or "", ["cost", "saving"]),
        "process_improvements": _impact_metric(role.description or "", ["process", "productivity"]),
        "technologies_used": _infer_skills(role.description or ""),
    }


def _join_location(city: str | None, state: str | None) -> str | None:
    parts = [part for part in [city, state] if part]
    return ", ".join(parts) if parts else None


def _infer_seniority_level(title: str | None, years: int | None) -> str | None:
    lowered = (title or "").lower()
    if any(token in lowered for token in ["chief", "cxo", "ceo", "cto", "cfo", "coo"]):
        return "C-Level"
    if "vp" in lowered or "vice president" in lowered:
        return "VP"
    if "director" in lowered:
        return "Director"
    if "manager" in lowered:
        return "Manager"
    if "lead" in lowered or "principal" in lowered or "staff" in lowered:
        return "Lead"
    if "senior" in lowered or (years is not None and years >= 5):
        return "Senior"
    if "associate" in lowered:
        return "Associate"
    if years is not None and years <= 1:
        return "Entry-Level"
    if years is not None and years < 5:
        return "Mid-Level"
    return None


def _infer_career_track(title: str | None, skills: list[str]) -> str | None:
    text = " ".join([title or "", *skills]).lower()
    if any(token in text for token in ["engineer", "developer", "python", "aws", "sql"]):
        return "Technical"
    if any(token in text for token in ["manager", "director", "lead", "head"]):
        return "Management"
    if any(token in text for token in ["sales", "business development"]):
        return "Sales"
    if any(token in text for token in ["marketing", "seo", "content"]):
        return "Marketing"
    if "operations" in text:
        return "Operations"
    if any(token in text for token in ["hr", "recruit", "talent acquisition"]):
        return "HR"
    if "finance" in text:
        return "Finance"
    return None


def _infer_function(title: str | None) -> str | None:
    track = _infer_career_track(title, [])
    return track


def _infer_primary_industry(text: str) -> str | None:
    lowered = text.lower()
    industries = {
        "Healthcare": ["healthcare", "hospital", "medical"],
        "Staffing & Recruiting": ["staffing", "recruiting", "talent acquisition"],
        "SaaS": ["saas", "software as a service"],
        "Finance": ["finance", "banking", "fintech"],
        "Manufacturing": ["manufacturing", "factory"],
        "Retail": ["retail", "ecommerce", "e-commerce"],
        "Education": ["education", "edtech", "school", "university"],
        "Government": ["government", "public sector"],
    }
    for industry, aliases in industries.items():
        if any(alias in lowered for alias in aliases):
            return industry
    return None


def _split_skill_groups(skills: list[str]) -> tuple[list[str], list[str]]:
    soft_aliases = {
        "Leadership",
        "Communication",
        "Negotiation",
        "Problem Solving",
        "Stakeholder Management",
    }
    hard: list[str] = []
    soft: list[str] = []
    for skill in skills:
        if skill in soft_aliases:
            soft.append(skill)
        else:
            hard.append(skill)
    return hard, soft


def _emerging_skills(skills: list[str]) -> list[str]:
    emerging = {"AI", "LLM", "LangGraph", "RAG", "MLOps", "Machine Learning", "Data Science"}
    return [skill for skill in skills if skill in emerging]


def _average_tenure(work_exp: list[WorkExperience]) -> float | None:
    if not work_exp:
        return None
    years = ResumeParserService._estimate_years_experience(work_exp)
    if years is None:
        return None
    return round(years / max(1, len(work_exp)), 1)


def _infer_promotions(work_exp: list[WorkExperience]) -> int:
    promotions = 0
    by_company: dict[str, int] = {}
    for role in work_exp:
        if not role.company:
            continue
        by_company[role.company.lower()] = by_company.get(role.company.lower(), 0) + 1
    for count in by_company.values():
        if count > 1:
            promotions += count - 1
    return promotions


def _career_stability_score(work_exp: list[WorkExperience]) -> int | None:
    avg = _average_tenure(work_exp)
    if avg is None:
        return None
    return max(1, min(100, round(avg / 3 * 100)))


def _infer_languages(text: str) -> list[str]:
    lowered = text.lower()
    languages = [
        "English",
        "Hindi",
        "Tamil",
        "Telugu",
        "Kannada",
        "Malayalam",
        "Marathi",
        "Bengali",
    ]
    return [lang for lang in languages if lang.lower() in lowered]


def _infer_certifications(text: str) -> list[dict]:
    certs: list[dict] = []
    patterns = [
        "AWS Certified Solutions Architect",
        "PMP",
        "SHRM",
        "Google Analytics",
        "HubSpot",
    ]
    lowered = text.lower()
    for cert in patterns:
        if cert.lower() in lowered:
            certs.append(
                {
                    "certification_name": cert,
                    "issuing_organization": None,
                    "issue_date": None,
                    "expiry_date": None,
                    "credential_id": None,
                    "verification_url": None,
                }
            )
    return certs


def _infer_awards(text: str) -> list[str]:
    return _lines_with_keywords(text, ["award", "honor", "recognition", "winner"])


def _infer_achievements(text: str) -> list[dict]:
    achievements = []
    keywords = ["increased", "reduced", "improved", "managed", "%", "$", "INR"]
    for line in _lines_with_keywords(text, keywords):
        achievements.append(
            {
                "description": line,
                "revenue_generated": _impact_metric(line, ["revenue", "sales"]),
                "revenue_influenced": None,
                "cost_reduction": _impact_metric(line, ["cost", "saving"]),
                "productivity_improvements": _impact_metric(line, ["productivity", "improved"]),
                "team_growth": _impact_metric(line, ["team", "managed"]),
                "market_expansion": _impact_metric(line, ["market", "expansion"]),
                "product_launches": _impact_metric(line, ["launch", "launched"]),
            }
        )
    return achievements[:20]


def _infer_domain_expertise(title: str | None, skills: list[str]) -> list[str]:
    text = " ".join([title or "", *skills]).lower()
    domains = {
        "Talent Acquisition": ["recruit", "talent acquisition"],
        "Sales Operations": ["salesforce", "sales operations"],
        "Product Management": ["product management", "product manager"],
        "Data Engineering": ["data engineering", "spark", "airflow"],
        "Cybersecurity": ["security", "cybersecurity"],
        "Customer Success": ["customer success"],
        "Marketing Automation": ["hubspot", "marketing automation"],
    }
    return [
        domain for domain, aliases in domains.items() if any(alias in text for alias in aliases)
    ]


def _infer_career_interests(text: str, skills: list[str]) -> list[str]:
    lowered = " ".join([text, *skills]).lower()
    interests = {
        "Leadership Aspirations": ["leadership", "manage", "manager"],
        "AI Interest": ["ai", "llm", "machine learning"],
        "Entrepreneurship": ["founder", "startup", "entrepreneur"],
        "Product Management Transition": ["product management", "product manager"],
    }
    return [
        name for name, aliases in interests.items() if any(alias in lowered for alias in aliases)
    ]


def _infer_transition_indicators(parsed: ParsedResume) -> dict:
    return {
        "possible_pivots": [],
        "signals": {
            "new_certifications": _infer_certifications(parsed.raw_text or ""),
            "side_projects": [],
            "recent_learning": _infer_learning_growth(parsed.raw_text or ""),
            "content_engagement": [],
        },
    }


def _infer_learning_growth(text: str) -> dict:
    return {
        "courses_completed": _lines_with_keywords(text, ["course", "coursera", "udemy"]),
        "recent_certifications": _infer_certifications(text),
        "new_skills_added": [],
        "books_mentioned": _lines_with_keywords(text, ["book"]),
        "workshops": _lines_with_keywords(text, ["workshop"]),
        "bootcamps": _lines_with_keywords(text, ["bootcamp"]),
    }


def _next_roles(title: str | None, track: str | None, seniority: str | None) -> list[str]:
    if track == "Technical":
        if seniority in {"Senior", "Lead"}:
            return ["Lead Engineer", "Staff Engineer", "Engineering Manager"]
        return ["Senior Software Engineer", "Backend Engineer", "Full Stack Engineer"]
    if track == "Sales":
        return ["Senior Account Executive", "Sales Manager", "Customer Success Manager"]
    if track == "HR":
        return ["Talent Acquisition Manager", "Recruitment Operations Lead", "HR Business Partner"]
    if track == "Management":
        return ["Senior Manager", "Director", "Head of Function"]
    return [f"Senior {title}" if title else "Senior Specialist"]


def _medium_term_roles(track: str | None) -> list[str]:
    return {
        "Technical": ["Engineering Manager", "Principal Engineer", "Director of Engineering"],
        "Sales": ["Sales Manager", "Regional Sales Lead", "Director of Sales"],
        "HR": ["Director of Talent Acquisition", "Head of Recruiting", "People Operations Lead"],
        "Management": ["Director", "Head of Department", "VP"],
    }.get(track or "", ["Manager", "Director", "Function Lead"])


def _long_term_roles(track: str | None) -> list[str]:
    return {
        "Technical": ["VP Engineering", "CTO", "Founder"],
        "Sales": ["VP Sales", "Chief Revenue Officer", "Founder"],
        "HR": ["Chief People Officer", "COO", "Founder"],
        "Management": ["VP", "COO", "Founder"],
    }.get(track or "", ["VP", "COO", "Founder"])


def _gap_analysis(track: str | None, skills: list[str]) -> dict:
    current = {skill.lower() for skill in skills}
    desired = {
        "Technical": ["system design", "cloud architecture", "leadership"],
        "Sales": ["crm", "forecasting", "enterprise sales"],
        "HR": ["people analytics", "workforce planning", "hr strategy"],
        "Management": ["strategy", "budget ownership", "people management"],
    }.get(track or "", ["leadership", "communication", "domain depth"])
    missing = [skill for skill in desired if skill.lower() not in current]
    return {
        "missing_skills": missing,
        "missing_certifications": [],
        "missing_experience": [],
    }


def _recommended_actions(gaps: dict) -> dict:
    missing_skills = gaps.get("missing_skills") or []
    return {
        "courses": [f"Take a focused course on {skill}" for skill in missing_skills[:3]],
        "certifications": gaps.get("missing_certifications") or [],
        "projects": [
            f"Build a portfolio project demonstrating {skill}" for skill in missing_skills[:2]
        ],
        "networking_activities": ["Connect with 5 leaders in target roles each month"],
        "leadership_opportunities": ["Volunteer to lead a cross-functional project"],
    }


def _employability_score(parsed: ParsedResume) -> int:
    score = 40
    if parsed.current_title:
        score += 15
    if parsed.years_experience:
        score += 15
    if parsed.skills:
        score += min(20, len(parsed.skills) * 2)
    if parsed.work_experience:
        score += 10
    return min(100, score)


def _market_competitiveness_score(parsed: ParsedResume) -> int:
    score = _employability_score(parsed)
    if _emerging_skills(parsed.skills):
        score += 10
    return min(100, score)


def _duration_label(start_date: str | None, end_date: str | None) -> str | None:
    if not start_date:
        return None
    return f"{start_date} - {end_date or 'Present'}"


def _sentences(text: str | None) -> list[str]:
    if not text:
        return []
    return [part.strip(" .") for part in re.split(r"[.;\n]", text) if part.strip()][:12]


def _lines_with_keywords(text: str, keywords: list[str]) -> list[str]:
    lines = [line.strip(" •-\t") for line in text.splitlines() if line.strip()]
    return [line for line in lines if any(keyword.lower() in line.lower() for keyword in keywords)][
        :20
    ]


def _impact_metric(text: str, keywords: list[str]) -> str | None:
    if not text or not any(keyword in text.lower() for keyword in keywords):
        return None
    match = re.search(r"(?:[$₹]\s?\d[\d,.]*[KkMmCrcr]*|\d+(?:\.\d+)?%)", text)
    return match.group(0) if match else None


def _has_any(text: str | None, keywords: list[str]) -> bool:
    lowered = (text or "").lower()
    return any(keyword.lower() in lowered for keyword in keywords)


def _first_int_after(text: str, keywords: list[str]) -> int | None:
    lowered = text.lower()
    if not any(keyword.lower() in lowered for keyword in keywords):
        return None
    match = re.search(r"\b(\d{1,4})\b", text)
    return int(match.group(1)) if match else None


def _year_from_date(value: str | None) -> int | None:
    if not value:
        return None
    match = re.search(r"\b(19|20)\d{2}\b", value)
    return int(match.group(0)) if match else None


# ── Lenient JSON / string helpers ───────────────────────────────────────────────


def _loads_json_lenient(content: str) -> object:
    """Parse JSON that may be wrapped in markdown fences or have leading prose."""
    if not content:
        return None
    text = content.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\n?", "", text)
        text = re.sub(r"\n?```$", "", text).strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        # Grab the largest {...} span and retry.
        start, end = text.find("{"), text.rfind("}")
        if 0 <= start < end:
            try:
                return json.loads(text[start : end + 1])
            except json.JSONDecodeError:
                return None
        return None


def _clean_str(value: object) -> str | None:
    if not isinstance(value, str):
        return None
    v = value.strip()
    return v or None


def _clean_int(value: object) -> int | None:
    """Coerce a numeric/clean string to a positive int; None otherwise."""
    if isinstance(value, bool):
        return None
    if isinstance(value, (int, float)):
        return int(value) if value >= 0 else None
    if isinstance(value, str):
        m = re.search(r"-?\d[\d,]*", value)
        if m:
            n = int(m.group(0).replace(",", ""))
            return n if n >= 0 else None
    return None


# ── Indian-résumé field extractors (CTC + notice period) ──────────────────────


def _infer_ctc(text: str) -> tuple[int | None, int | None, int | None]:
    """(expected_min, expected_max, current) in INR p.a. from labelled lines like
    'Expected CTC: 25 LPA', 'Current CTC: 18 lpa', 'CTC: 12-15 LPA'."""

    def _to_inr(num: str, unit: str) -> int:
        val = float(num)
        u = unit.lower()
        if "cr" in u:
            return int(val * 10_000_000)
        if "lpa" in u or "lakh" in u or "lac" in u or u == "l":
            return int(val * 100_000)
        # Bare number ≥ 1000 → already absolute INR; small number → assume LPA.
        return int(val) if val >= 1000 else int(val * 100_000)

    unit_re = r"(lpa|lakhs?|lacs?|cr|crores?|l)\b"
    exp_min = exp_max = cur = None
    for label, kind in (("current", "cur"), ("expected", "exp"), ("ctc", "exp")):
        # range: "<label> ... 12-15 LPA"
        rng = re.search(
            rf"{label}[^\n]*?(\d+(?:\.\d+)?)\s*[-–to]+\s*(\d+(?:\.\d+)?)\s*{unit_re}",
            text,
            re.IGNORECASE,
        )
        if rng:
            lo = _to_inr(rng.group(1), rng.group(3))
            hi = _to_inr(rng.group(2), rng.group(3))
            if kind == "exp" and exp_min is None:
                exp_min, exp_max = lo, hi
            elif kind == "cur" and cur is None:
                cur = lo
            continue
        single = re.search(rf"{label}[^\n]*?(\d+(?:\.\d+)?)\s*{unit_re}", text, re.IGNORECASE)
        if single:
            amt = _to_inr(single.group(1), single.group(2))
            if kind == "exp" and exp_min is None:
                exp_min = amt
            elif kind == "cur" and cur is None:
                cur = amt
    return exp_min, exp_max, cur


def _infer_notice_period(text: str) -> int | None:
    """Days from 'Notice Period: 30 days' / '2 months' / 'Immediate'/'Serving'."""
    m = re.search(r"notice\s*period[^\n]{0,40}", text, re.IGNORECASE)
    seg = m.group(0).lower() if m else ""
    if not seg:
        return None
    if "immediate" in seg:
        return 0
    months = re.search(r"(\d+)\s*month", seg)
    if months:
        return int(months.group(1)) * 30
    days = re.search(r"(\d+)\s*day", seg)
    if days:
        return int(days.group(1))
    return None


def _normalise_parsed_resume(parsed: ParsedResume, *, source: str) -> ParsedResume:
    """Canonicalize parser output from any tier and attach quality metadata."""
    if parsed.raw_text:
        parsed.phone = _normalise_india_phone(parsed.phone) or _infer_india_phone(parsed.raw_text)
        parsed.linkedin_url = _normalise_profile_url(
            parsed.linkedin_url, "linkedin"
        ) or _infer_profile_url(
            parsed.raw_text,
            "linkedin",
        )
        parsed.github_url = _normalise_profile_url(
            parsed.github_url, "github"
        ) or _infer_profile_url(
            parsed.raw_text,
            "github",
        )
        if not parsed.location_city or not parsed.location_state:
            city, state = _infer_location(parsed.raw_text)
            parsed.location_city = parsed.location_city or city
            parsed.location_state = parsed.location_state or state
    else:
        parsed.phone = _normalise_india_phone(parsed.phone)
        parsed.linkedin_url = _normalise_profile_url(parsed.linkedin_url, "linkedin")
        parsed.github_url = _normalise_profile_url(parsed.github_url, "github")

    parsed.location_city, parsed.location_state = _normalise_location_pair(
        parsed.location_city,
        parsed.location_state,
    )
    parsed.skills = _normalise_skill_list(parsed.skills)

    # #1: title quality gate. A LinkedIn headline/tagline ("Helping recruiters …
    # | GTM Lead") is not a job title — it pollutes matching's title-affinity and
    # recruiter search. When the current_title looks like a tagline, demote it to
    # the headline and use the real job title from work history instead. If there
    # is no work title, null current_title so it surfaces as a profile gap to
    # confirm rather than a confident-but-wrong value.
    if _looks_like_tagline(parsed.current_title):
        if not (parsed.headline or "").strip():
            parsed.headline = parsed.current_title
        work_title = _first_work_title(parsed.work_experience)
        parsed.current_title = work_title  # may be None → treated as a gap

    if parsed.years_experience is None and parsed.work_experience:
        parsed.years_experience = _estimate_years_static(parsed.work_experience)

    # #22: normalise free-text dates to ISO YYYY-MM and trust the MATH over the
    # stated number when work history is dated. People round up ("10+ years");
    # interval-merged tenure from actual roles is the honest figure. The stated
    # value wins only when history is undated or the two roughly agree.
    for exp in parsed.work_experience or []:
        start = parse_resume_date(exp.start_date)
        end = parse_resume_date(exp.end_date)
        if start:
            exp.start_date = start.strftime("%Y-%m")
        if end:
            exp.end_date = end.strftime("%Y-%m")
        elif exp.end_date and any(w in str(exp.end_date).lower() for w in _PRESENT_WORDS):
            exp.end_date = None
            exp.is_current = True
    computed = compute_tenure_years(parsed.work_experience or [])
    if computed is not None:
        stated = parsed.years_experience
        if stated is None or abs(int(stated) - computed) > 3:
            parsed.years_experience = computed

    if parsed.years_experience is not None:
        parsed.years_experience = max(0, min(60, int(parsed.years_experience)))

    parsed.parser_metadata = {
        **(parsed.parser_metadata or {}),
        "source": parsed.parser_metadata.get("source") or source,
        "quality_score": _resume_quality_score(parsed),
        "normalized_fields": _normalised_field_names(parsed),
    }
    return parsed


def _normalised_field_names(parsed: ParsedResume) -> list[str]:
    field_names = [
        "full_name",
        "email",
        "phone",
        "headline",
        "summary",
        "current_title",
        "current_company",
        "years_experience",
        "skills",
        "work_experience",
        "education",
        "linkedin_url",
        "github_url",
        "location_city",
        "location_state",
    ]
    return [field for field in field_names if bool(getattr(parsed, field))]


def _resume_quality_score(parsed: ParsedResume) -> int:
    score = 0
    weights = {
        "full_name": 8,
        "email": 8,
        "phone": 8,
        "current_title": 12,
        "current_company": 10,
        "years_experience": 10,
        "location_city": 6,
        "location_state": 4,
        "linkedin_url": 6,
        "github_url": 4,
        "summary": 6,
    }
    for field, weight in weights.items():
        if getattr(parsed, field):
            score += weight
    score += min(14, len(parsed.skills) * 2)
    score += min(8, len(parsed.work_experience) * 4)
    score += min(4, len(parsed.education) * 2)
    return min(100, score)


def _normalise_india_phone(value: str | None) -> str | None:
    if not value:
        return None
    digits = re.sub(r"\D", "", value)
    if len(digits) == 10 and digits[0] in "6789":
        return f"+91{digits}"
    if len(digits) == 11 and digits.startswith("0") and digits[1] in "6789":
        return f"+91{digits[1:]}"
    if len(digits) == 12 and digits.startswith("91") and digits[2] in "6789":
        return f"+{digits}"
    return None


def _infer_india_phone(text: str) -> str | None:
    for match in re.finditer(
        r"(?:\+?91[\s-]*)?(?:0[\s-]*)?[6-9](?:[\s-]*\d){9}",
        text,
    ):
        phone = _normalise_india_phone(match.group(0))
        if phone:
            return phone
    return None


def _infer_profile_url(text: str, kind: str) -> str | None:
    if kind == "linkedin":
        pattern = r"(?:https?://)?(?:www\.)?linkedin\.com/(?:in|pub)/[^\s,;)]*"
    else:
        pattern = r"(?:https?://)?(?:www\.)?github\.com/[^\s,;)]*"
    match = re.search(pattern, text, flags=re.IGNORECASE)
    return _normalise_profile_url(match.group(0), kind) if match else None


def _normalise_profile_url(value: str | None, kind: str) -> str | None:
    if not value:
        return None
    raw = value.strip().strip(".,;:)]}>")
    if not raw:
        return None
    if not raw.startswith(("http://", "https://")):
        raw = f"https://{raw}"
    raw = raw.replace("http://", "https://", 1)
    raw = re.sub(r"/+$", "", raw)
    lower = raw.lower()
    if kind == "linkedin":
        if "linkedin.com/" not in lower or not ("/in/" in lower or "/pub/" in lower):
            return None
        raw = re.sub(
            r"https://(?:[a-z]{2,3}\.)?linkedin\.com",
            "https://www.linkedin.com",
            raw,
            flags=re.IGNORECASE,
        )
    elif kind == "github" and "github.com/" not in lower:
        return None
    return raw


def _normalise_location_pair(
    city: str | None,
    state: str | None,
) -> tuple[str | None, str | None]:
    clean_city = _clean_str(city)
    clean_state = _clean_str(state)
    city_key = clean_city.casefold() if clean_city else ""
    state_key = clean_state.casefold() if clean_state else ""

    display_city = _CITY_DISPLAY_OVERRIDES.get(city_key, clean_city)
    if display_city and display_city.casefold() in _INDIAN_LOCATIONS:
        display_city = _CITY_DISPLAY_OVERRIDES.get(
            display_city.casefold(),
            display_city.title(),
        )
        clean_state = clean_state or _INDIAN_LOCATIONS[display_city.casefold()]
    if state_key:
        clean_state = _INDIAN_STATE_ALIASES.get(state_key, clean_state)
    return display_city, clean_state


# Section/CV-furniture words that masquerade as skills when a parser splits a
# résumé's "Languages / Interests / Declaration" blocks. Distinct from
# _SECTION_HEADERS (which guards name detection).
_SKILL_SECTION_NOISE = frozenset(
    {
        "languages",
        "language",
        "interests",
        "hobbies",
        "references",
        "achievements",
        "awards",
        "declaration",
        "personal details",
        "personal information",
        "strengths",
        "activities",
    }
)

# A skill phrase that opens with one of these is almost always a sentence
# fragment ("i personally:", "used by clients", "managed the team"), not a skill.
_SKILL_FRAGMENT_PREFIXES = (
    "i ",
    "we ",
    "my ",
    "our ",
    "the ",
    "a ",
    "an ",
    "used ",
    "using ",
    "based ",
    "including ",
    "such as",
    "responsible ",
    "worked ",
    "managed ",
    "branded",
    "e.g",
    "eg ",
    "etc",
)


def _is_low_quality_skill(key: str) -> bool:
    """
    True for resume-parse artifacts that aren't real skills — the noise you see
    when a parser over-splits ("i personally:", "used by", "languages",
    "english (professional working)"). Conservative: real multi-word skills like
    "product management" or "stakeholder communication" pass.
    """
    if key.endswith(":") or ":" in key:
        return True
    if key in _SKILL_SECTION_NOISE:
        return True
    if any(key.startswith(p) for p in _SKILL_FRAGMENT_PREFIXES):
        return True
    # Skills are short noun phrases; >4 words is a sentence fragment.
    if len(key.split()) > 4:
        return True
    # Pure language-proficiency entries ("english (professional working)").
    if re.match(r"^[a-z]+\s*\((?:native|fluent|professional|basic|working)", key):
        return True
    return False


def _canonical_skill(value: str) -> str | None:
    key = value.strip(" .;|·-").casefold()
    key = re.sub(r"\s+", " ", key)
    if not key or key in _SECTION_HEADERS or key in _JUNK_SKILLS:
        return None
    if len(key) <= 1 or len(key) > 40:
        return None
    # The bundled vocabulary is the whitelist: a recognised skill is always kept
    # and returned with its canonical display label ("postgres" → "PostgreSQL"),
    # bypassing the heuristic junk filter (which can over-trim real skills).
    if _vocab_known(key):
        return _vocab_display(key)
    if _is_low_quality_skill(key):
        return None
    for canonical, aliases in _SKILL_ALIASES.items():
        if key == canonical.casefold():
            return canonical
        for alias in aliases:
            if key == alias.casefold():
                return canonical
    return key


def _normalise_skill_list(raw: object) -> list[str]:
    if not isinstance(raw, list):
        return []
    out: list[str] = []
    seen: set[str] = set()
    for item in raw:
        if not isinstance(item, str):
            continue
        canonical = _canonical_skill(item)
        if not canonical or canonical in seen:
            continue
        seen.add(canonical)
        out.append(canonical)
    return out[:40]


def _clean_skill_list(raw: object) -> list[str]:
    if not isinstance(raw, list):
        return []
    out: list[str] = []
    for item in raw:
        if not isinstance(item, str):
            continue
        s = item.strip(" .;|·-").lower()
        if 1 < len(s) <= 40 and s not in out:
            out.append(s)
    return _normalise_skill_list(out)


# ── Regex helpers ──────────────────────────────────────────────────────────────


def _normalise_text(text: str) -> str:
    return "\n".join(line.strip() for line in (text or "").replace("\r", "\n").splitlines())


def _first_match(text: str, pattern: str) -> str | None:
    match = re.search(pattern, text, flags=re.IGNORECASE)
    return match.group(0).strip() if match else None


def _looks_like_contact_line(line: str) -> bool:
    lowered = line.lower()
    contact_tokens = ("@", "linkedin.com", "github.com", "+91", "http")
    return any(token in lowered for token in contact_tokens) or bool(re.search(r"\d{6,}", line))


def _infer_name(lines: list[str]) -> str | None:
    """Pick the most name-like line from the top of the resume."""
    for line in lines[:10]:
        lowered = line.lower().strip(" :")
        if lowered in _SECTION_HEADERS:
            continue
        if _looks_like_contact_line(line):
            continue
        words = line.split()
        # A name is short, mostly alphabetic, 1-4 tokens, no digits.
        if not (1 <= len(words) <= 4):
            continue
        if any(ch.isdigit() for ch in line):
            continue
        alpha = re.sub(r"[^A-Za-z ]", "", line)
        if len(alpha) < 3:
            continue
        # Title-case or ALL CAPS names both accepted.
        if line.isupper() or all(w[:1].isupper() for w in words if w):
            return _titleize(line)
    # Fallback: first non-header, non-contact line.
    for line in lines[:8]:
        if line.lower().strip(" :") not in _SECTION_HEADERS and not _looks_like_contact_line(line):
            return line
    return None


def _titleize(line: str) -> str:
    return " ".join(w.capitalize() if w.isupper() else w for w in line.split())


_TAGLINE_OPENERS = (
    "helping",
    "building",
    "empowering",
    "passionate",
    "driving",
    "transforming",
    "enabling",
    "ex-",
    "ex ",
)


def _looks_like_tagline(title: str | None) -> bool:
    """
    True when a 'title' is really a LinkedIn headline / marketing tagline rather
    than a job title. Conservative: real titles like "Senior Backend Engineer"
    or "GTM Lead - AI Resume Builder" pass; "Helping recruiters … | GTM Lead"
    (pipes, marketing opener, or very long) does not.
    """
    t = (title or "").strip()
    if not t:
        return False
    low = t.lower()
    if "|" in t:
        return True
    if any(low.startswith(p) for p in _TAGLINE_OPENERS):
        return True
    if len(t) > 70 or len(t.split()) > 9:
        return True
    return False


def _first_work_title(work: list[WorkExperience]) -> str | None:
    """First non-empty job title, preferring the current role.

    Populates current_title even when the parser left the 'current' entry's title
    blank (e.g. a "Founder, LimeDock" line landing in company), by scanning the
    rest of the work history. A populated current_title also feeds the matching
    engine's title-affinity signal.
    """
    if not work:
        return None
    ordered = [w for w in work if getattr(w, "is_current", False)] + [
        w for w in work if not getattr(w, "is_current", False)
    ]
    for w in ordered:
        title = _clean_str(getattr(w, "title", None))
        if title:
            return title
    return None


def _infer_current_role(lines: list[str]) -> tuple[str | None, str | None]:
    sep = r"\s*(?:at|@|[-–—|,·])\s*"
    for line in lines[:25]:
        if _looks_like_contact_line(line):
            continue
        m = re.search(
            rf"(?P<title>[A-Za-z][A-Za-z0-9 /&.+-]{{2,80}}){sep}"
            rf"(?P<company>[A-Za-z][A-Za-z0-9 &.+'-]{{1,80}})",
            line,
        )
        if m:
            title = m.group("title").strip(" -–—|,·")
            company = m.group("company").strip(" -–—|,·")
            if _looks_like_role_title(title):
                return title, company
    role_keywords = (
        "engineer",
        "developer",
        "manager",
        "designer",
        "analyst",
        "consultant",
        "lead",
        "architect",
        "scientist",
        "specialist",
        "executive",
        "director",
        "intern",
        "associate",
        "officer",
        "head",
    )
    for line in lines[:15]:
        low = line.lower()
        if any(word in low for word in role_keywords) and not _looks_like_contact_line(line):
            return line.strip(" -–—|,·"), None
    return None, None


def _looks_like_role_title(text: str) -> bool:
    role_keywords = (
        "engineer",
        "developer",
        "manager",
        "designer",
        "analyst",
        "consultant",
        "lead",
        "architect",
        "scientist",
        "specialist",
        "executive",
        "director",
        "intern",
        "associate",
        "officer",
        "head",
        "senior",
        "junior",
        "founder",
    )
    return any(k in text.lower() for k in role_keywords)


def _infer_years_experience(text: str) -> int | None:
    patterns = [
        r"(\d{1,2})\+?\s*(?:years|yrs)\.?\s*(?:of\s+)?(?:experience|exp)",
        r"experience\s*(?:of|:)?\s*(\d{1,2})\+?\s*(?:years|yrs)",
        r"(\d{1,2})\+\s*(?:years|yrs)",
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            return min(60, int(m.group(1)))
    return None


def _estimate_years_static(work_exp: list[WorkExperience]) -> int | None:
    return ResumeParserService._estimate_years_experience(work_exp) if work_exp else None


# ── Date normalization + tenure math (#22) ─────────────────────────────────────

_MONTHS = {
    m: i + 1
    for i, names in enumerate(
        [
            ("jan", "january"),
            ("feb", "february"),
            ("mar", "march"),
            ("apr", "april"),
            ("may",),
            ("jun", "june"),
            ("jul", "july"),
            ("aug", "august"),
            ("sep", "sept", "september"),
            ("oct", "october"),
            ("nov", "november"),
            ("dec", "december"),
        ]
    )
    for m in names
}

_PRESENT_WORDS = ("present", "current", "now", "till date", "to date", "ongoing", "today")


def parse_resume_date(raw: str | None) -> date | None:
    """
    Parse the date formats people actually write on CVs into a real date
    (day pinned to 1): "Jan 2021", "January 2021", "Mar'21", "03/2021",
    "2021-03", "2021". "Present"-type words return None (caller treats a
    missing END date as today).
    """
    if not raw:
        return None
    s = str(raw).strip().lower().replace("’", "'")
    if any(w in s for w in _PRESENT_WORDS):
        return None
    # ISO-ish: 2021-03[-15]
    m = re.match(r"^(\d{4})[-/](\d{1,2})", s)
    if m:
        y, mo = int(m.group(1)), int(m.group(2))
        if 1 <= mo <= 12 and 1950 <= y <= 2100:
            return date(y, mo, 1)
    # Numeric: 03/2021 or 3-2021
    m = re.match(r"^(\d{1,2})[-/](\d{4})$", s)
    if m:
        mo, y = int(m.group(1)), int(m.group(2))
        if 1 <= mo <= 12 and 1950 <= y <= 2100:
            return date(y, mo, 1)
    # Month name + year: "jan 2021", "january, 2021", "mar'21"
    m = re.match(r"^([a-z]+)\.?,?\s*'?(\d{2,4})$", s)
    if m and m.group(1) in _MONTHS:
        y = int(m.group(2))
        y = y + 2000 if y < 100 else y  # '21 → 2021 (CVs don't reference the 1900s this way)
        if 1950 <= y <= 2100:
            return date(y, _MONTHS[m.group(1)], 1)
    # Bare year: "2021" → January of that year (conservative)
    m = re.match(r"^(\d{4})$", s)
    if m and 1950 <= int(m.group(1)) <= 2100:
        return date(int(m.group(1)), 1, 1)
    return None


def compute_tenure_years(work_exp: list[WorkExperience]) -> int | None:
    """
    Total experience from work history with OVERLAP MERGING — concurrent roles
    (freelance + day job) must not double-count. Returns None when no entry has
    a parseable start date.
    """
    today = date.today()
    intervals: list[tuple[date, date]] = []
    for exp in work_exp or []:
        start = parse_resume_date(exp.start_date)
        if not start:
            continue
        end = parse_resume_date(exp.end_date) or today
        if exp.is_current:
            end = today
        if end < start:
            start, end = end, start
        intervals.append((start, min(end, today)))
    if not intervals:
        return None
    intervals.sort()
    merged: list[tuple[date, date]] = [intervals[0]]
    for start, end in intervals[1:]:
        last_start, last_end = merged[-1]
        if start <= last_end:
            merged[-1] = (last_start, max(last_end, end))
        else:
            merged.append((start, end))
    months = sum((e.year - s.year) * 12 + (e.month - s.month) for s, e in merged)
    return max(0, round(months / 12))


def _infer_skills(text: str) -> list[str]:
    lowered = text.lower()
    explicit: list[str] = []
    skills_match = re.search(
        r"(?:technical skills|core competencies|key skills|skills|technologies)\s*[:\n]\s*"
        r"(?P<skills>.+?)"
        r"(?:\n\s*\n|work experience|professional experience|experience|education|"
        r"projects|certification|$)",
        text,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if skills_match:
        explicit = [
            item.strip(" .;|·-")
            for item in re.split(r"[,/•|\n\t]", skills_match.group("skills"))
            if item.strip(" .;|·-")
        ]
        explicit = [s for s in explicit if 1 < len(s) <= 40]

    detected: list[str] = []
    for canonical, aliases in _SKILL_ALIASES.items():
        for alias in aliases:
            if re.search(rf"(?<![\w]){alias}(?![\w])", lowered):
                detected.append(_display_skill(canonical))
                break

    # Map any explicit free-text skill onto a canonical label when we recognise
    # it, so "reactjs" and "React" don't both show up.
    normalised_explicit: list[str] = []
    for raw in explicit:
        key = raw.lower().strip()
        if key in _SECTION_HEADERS or key in {"work", "experience", "education", "projects"}:
            continue  # stray section header that slipped past the section boundary
        canonical = None
        for cand, aliases in _SKILL_ALIASES.items():
            if key == cand or key in aliases:
                canonical = cand
                break
        normalised_explicit.append(_display_skill(canonical) if canonical else raw)

    merged: list[str] = []
    seen: set[str] = set()
    for skill in [*detected, *normalised_explicit]:  # canonical first for cleaner labels
        low = skill.lower()
        if 1 < len(skill) <= 40 and low not in seen:
            seen.add(low)
            merged.append(skill)
    return merged[:30]


def _infer_location(text: str) -> tuple[str | None, str | None]:
    lowered = text.lower()
    display_overrides = {
        "bangalore": "Bengaluru",
        "cochin": "Kochi",
        "trivandrum": "Thiruvananthapuram",
    }
    # Longest city names first so "navi mumbai" beats "mumbai".
    for city in sorted(_INDIAN_LOCATIONS, key=len, reverse=True):
        if re.search(rf"\b{re.escape(city)}\b", lowered):
            display = display_overrides.get(city, city.title())
            return display, _INDIAN_LOCATIONS[city]
    return None, None


def _infer_work_experience(lines: list[str]) -> list[WorkExperience]:
    """
    Extract multiple work-experience entries by anchoring on date ranges.

    Heuristic: for each line containing a date range, the title/company usually
    sit on that line or the line just above it. Conservative — only emits entries
    we can attach a date range to, so we never fabricate roles.
    """
    _edu_kw = re.compile(
        r"\b(b\.?tech|btech|b\.?e\.?|bachelor|m\.?tech|mtech|m\.?e\.?|master|mba|bca|mca|"
        r"b\.?sc|m\.?sc|b\.?com|m\.?com|phd|diploma|university|college|institute of|"
        r"\biit\b|\bnit\b|\biiit\b|\biim\b)\b",
        re.IGNORECASE,
    )
    entries: list[WorkExperience] = []
    in_education = False
    for i, line in enumerate(lines):
        header = line.lower().strip(" :")
        if header in _SECTION_HEADERS:
            in_education = "education" in header or "academic" in header
            continue
        m = _DATE_RANGE.search(line)
        if not m:
            continue
        # Don't mistake degree lines ("B.Tech … 2014 - 2018") for jobs.
        if in_education or _edu_kw.search(line):
            continue
        start_raw, end_raw = m.group(1), m.group(2)
        is_current = bool(
            re.search(r"present|current|now|till date|ongoing", end_raw, re.IGNORECASE)
        )

        # Text on the date line minus the dates, plus the line above, are candidates
        # for "Title at Company" / "Title, Company".
        same_line = _DATE_RANGE.sub("", line).strip(" -–—|,·")
        above = lines[i - 1].strip() if i > 0 else ""
        title, company = None, None
        for candidate in (same_line, above):
            if candidate and not _looks_like_contact_line(candidate):
                t, c = _split_title_company(candidate)
                title = title or t
                company = company or c
            if title and company:
                break

        if title or company:
            entries.append(
                WorkExperience(
                    company=company,
                    title=title,
                    start_date=_norm_date(start_raw),
                    end_date=None if is_current else _norm_date(end_raw),
                    is_current=is_current,
                )
            )
        if len(entries) >= 8:
            break
    return entries


def _split_title_company(text: str) -> tuple[str | None, str | None]:
    m = re.search(
        r"(?P<title>[A-Za-z][A-Za-z0-9 /&.+-]{2,80})\s*(?:at|@|[-–—|,·])\s*"
        r"(?P<company>[A-Za-z][A-Za-z0-9 &.+'-]{1,80})",
        text,
    )
    if m:
        return m.group("title").strip(" -–—|,·"), m.group("company").strip(" -–—|,·")
    if _looks_like_role_title(text):
        return text.strip(" -–—|,·"), None
    return None, text.strip(" -–—|,·") or None


def _norm_date(raw: str) -> str | None:
    """Normalise a fuzzy date token to 'YYYY-MM' or 'YYYY' where possible."""
    raw = raw.strip()
    months = {
        "jan": "01",
        "feb": "02",
        "mar": "03",
        "apr": "04",
        "may": "05",
        "jun": "06",
        "jul": "07",
        "aug": "08",
        "sep": "09",
        "oct": "10",
        "nov": "11",
        "dec": "12",
    }
    m = re.match(r"([a-zA-Z]{3})[a-z]*\.?\s*'?(\d{2,4})", raw)
    if m:
        mon = months.get(m.group(1).lower())
        year = m.group(2)
        if len(year) == 2:
            year = ("20" if int(year) < 50 else "19") + year
        return f"{year}-{mon}" if mon else year
    m = re.match(r"(\d{1,2})/(\d{4})", raw)
    if m:
        return f"{m.group(2)}-{int(m.group(1)):02d}"
    m = re.match(r"(\d{4})", raw)
    if m:
        return m.group(1)
    return None


def _infer_education(lines: list[str]) -> list[Education]:
    # Strong, unambiguous degree tokens (case-insensitive). These never collide
    # with ordinary English words, so a match alone is a reliable signal.
    strong_degree_kw = re.compile(
        r"\b(b\.?tech|btech|m\.?tech|mtech|bachelor|master|mba|bca|mca|"
        r"b\.?sc|m\.?sc|b\.?com|m\.?com|ph\.?d|diploma)\b",
        re.IGNORECASE,
    )
    # Ambiguous 2-letter abbreviations (B.E./M.E./B.A./M.A.) collide with the
    # words "be", "me", "ba", "ma". Match ONLY the uppercase/dotted form so a
    # line like "About me" or "I'll be there" can't masquerade as a degree.
    weak_degree_kw = re.compile(r"\b([BM]\.?[EA])\.?\b")
    inst_kw = re.compile(
        r"\b(university|college|institute|iit|nit|iiit|iim|school|academy|polytechnic)\b",
        re.IGNORECASE,
    )
    out: list[Education] = []
    seen: set[str] = set()
    in_education = False
    for line in lines:
        header = line.lower().strip(" :")
        if header in _SECTION_HEADERS:
            in_education = "education" in header or "academic" in header
            continue
        if _looks_like_contact_line(line):
            continue

        strong = strong_degree_kw.search(line)
        weak = weak_degree_kw.search(line)
        has_inst = inst_kw.search(line)
        years = re.findall(r"(?:19|20)\d{2}", line)
        year = years[-1] if years else None  # graduation year = latest year on line

        # Acceptance gate — reject lines whose only signal is noise:
        #  • institution keyword → accept only inside the education section, or
        #    when a year/degree co-occurs (avoids grabbing "Institute" in a
        #    company name from the experience section).
        #  • strong degree → accept inside the section, or with a year alongside.
        #  • weak degree (B.E./M.E./…) → only inside the education section, since
        #    even uppercase "BE"/"MA" is too noisy to trust on its own elsewhere.
        if has_inst:
            accept = in_education or bool(year) or bool(strong) or bool(weak)
        elif strong:
            accept = in_education or bool(year)
        elif weak:
            accept = in_education
        else:
            accept = False
        if not accept:
            continue

        degree = None
        if strong:
            degree = strong.group(0).strip()
        elif weak:
            degree = weak.group(0).strip().upper()

        institution = None
        if has_inst:
            # Pull just the comma/pipe segment that contains the institution
            # keyword (e.g. "IIT Delhi") rather than the whole line, dropping any
            # trailing year tokens.
            for seg in re.split(r"[,|•·]", line):
                if inst_kw.search(seg):
                    institution = re.sub(r"\b(?:19|20)\d{2}\b.*$", "", seg).strip(" -–—.")
                    break
            if not institution:
                institution = re.sub(r"\b(?:19|20)\d{2}\b.*$", "", line).strip(" -–—.")
            # Reject prose masquerading as an institution: a real school name
            # doesn't contain a degree token and isn't a long sentence. (Drops
            # cases like "Completed MBA in" extracted from a paragraph.)
            if institution and (
                strong_degree_kw.search(institution)
                or weak_degree_kw.search(institution)
                or len(institution.split()) > 8
            ):
                institution = None

        # A surviving entry must carry at least an institution or a degree.
        if not (institution or degree):
            continue
        key = f"{(institution or '').casefold()}|{(degree or '').casefold()}"
        if key in seen:
            continue
        seen.add(key)
        out.append(
            Education(
                institution=institution,
                degree=degree,
                end_date=year,
            )
        )
        if len(out) >= 6:
            break
    return out


def _infer_summary(
    lines: list[str],
    title: str | None,
    company: str | None,
    years: int | None,
    skills: list[str],
) -> str | None:
    # Prefer an explicit Summary/Objective section if present.
    text = "\n".join(lines)
    m = re.search(
        r"(?:professional summary|summary|profile|objective|about me|about)\s*[:\n]\s*"
        r"(?P<body>.+?)(?:\n\s*\n|experience|education|skills|$)",
        text,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if m:
        body = re.sub(r"\s+", " ", m.group("body")).strip()
        if 20 <= len(body) <= 600:
            return body
    # Otherwise synthesise a short one.
    parts = [
        f"{title} at {company}" if title and company else title,
        f"{years} years of experience" if years else None,
        f"Skills: {', '.join(skills[:8])}" if skills else None,
    ]
    parts = [p for p in parts if p]
    return ". ".join(parts) if parts else None
