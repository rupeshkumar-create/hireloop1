"""Export resume HTML fragments to Word (.docx) for candidate downloads."""

from __future__ import annotations

import io
import re
from html.parser import HTMLParser
from urllib.parse import quote

from docx import Document
from docx.shared import Pt


def ascii_download_filename(name: str, *, fallback: str = "resume", ext: str = "html") -> str:
    """Latin-1-safe filename for Content-Disposition headers."""
    base = (name or fallback).replace("/", "-")
    for src, dst in (("—", "-"), ("–", "-"), ("'", ""), ('"', "")):
        base = base.replace(src, dst)
    slug = re.sub(r"[^\w.\- ]+", "", base, flags=re.ASCII).strip(" .-_")
    if not slug:
        slug = fallback
    slug = slug[:80]
    return f"{slug}.{ext}" if ext else slug


def content_disposition_header(disposition: str, basename: str, *, ext: str = "html") -> str:
    """Build a Content-Disposition value safe for HTTP headers."""
    safe = ascii_download_filename(basename, ext=ext)
    full = f"{basename}.{ext}"
    encoded = quote(full)
    return f'{disposition}; filename="{safe}"; filename*=UTF-8\'\'{encoded}'


def _extract_body_html(html: str) -> str:
    inner = (html or "").strip()
    body_match = re.search(r"<body[^>]*>(.*)</body>", inner, re.IGNORECASE | re.DOTALL)
    if body_match:
        inner = body_match.group(1).strip()
    else:
        inner = re.sub(r"<head[^>]*>.*?</head>", "", inner, flags=re.IGNORECASE | re.DOTALL)
        inner = re.sub(r"</?(?:html|body)[^>]*>", "", inner, flags=re.IGNORECASE)
    # Drop toolbar / scripts from wrapped print documents.
    inner = re.sub(r'<div[^>]*class="[^"]*toolbar[^"]*"[^>]*>.*?</div>', "", inner, flags=re.DOTALL)
    inner = re.sub(r"<script[^>]*>.*?</script>", "", inner, flags=re.IGNORECASE | re.DOTALL)
    return inner.strip()


class _ResumeHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.blocks: list[tuple[str, str]] = []
        self._current_list: list[str] = []
        self._in_li = False
        self._li_parts: list[str] = []
        self._capture_tag: str | None = None
        self._capture_parts: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        t = tag.lower()
        if t in ("h1", "h2", "h3", "p"):
            self._flush_li()
            self._flush_capture()
            self._capture_tag = t
            self._capture_parts = []
        elif t == "ul":
            self._flush_capture()
            self._current_list = []
        elif t == "li":
            self._flush_capture()
            self._in_li = True
            self._li_parts = []
        elif t == "br" and self._capture_tag:
            self._capture_parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        t = tag.lower()
        if t in ("h1", "h2", "h3", "p") and self._capture_tag == t:
            text = "".join(self._capture_parts).strip()
            if text:
                self.blocks.append((t, text))
            self._capture_tag = None
            self._capture_parts = []
        elif t == "li":
            self._flush_li()
        elif t == "ul":
            if self._current_list:
                self.blocks.append(("ul", "\n".join(self._current_list)))
                self._current_list = []

    def handle_data(self, data: str) -> None:
        if self._in_li:
            self._li_parts.append(data)
        elif self._capture_tag:
            self._capture_parts.append(data)

    def _flush_li(self) -> None:
        if not self._in_li:
            return
        text = "".join(self._li_parts).strip()
        if text:
            self._current_list.append(text)
        self._in_li = False
        self._li_parts = []

    def _flush_capture(self) -> None:
        if self._capture_tag and self._capture_parts:
            text = "".join(self._capture_parts).strip()
            if text:
                self.blocks.append((self._capture_tag, text))
        self._capture_tag = None
        self._capture_parts = []

    def close(self) -> None:
        self._flush_li()
        self._flush_capture()
        if self._current_list:
            self.blocks.append(("ul", "\n".join(self._current_list)))
        super().close()


def html_resume_to_docx(body_html: str, *, title: str = "Resume") -> bytes:
    """Convert resume HTML fragment (or wrapped print doc) to a .docx byte stream."""
    parser = _ResumeHTMLParser()
    parser.feed(_extract_body_html(body_html))
    parser.close()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    if not parser.blocks:
        doc.add_paragraph(_strip_tags(body_html)[:8000] or title)
    else:
        for kind, text in parser.blocks:
            _add_block(doc, kind, text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _strip_tags(html: str) -> str:
    return re.sub(r"<[^>]+>", " ", html or "").replace("&nbsp;", " ").strip()


def _add_block(doc: Document, kind: str, text: str) -> None:
    if kind == "h1":
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(20)
        p.paragraph_format.space_after = Pt(4)
    elif kind == "h2":
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    elif kind == "h3":
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(6)
    elif kind == "ul":
        for line in text.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip(), style="List Bullet")
    else:
        doc.add_paragraph(text)
